from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Solverstreet_Technologies_BA_Project_Documentation.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 103, 124)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D7DBE2"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths):
            if index >= len(row.cells):
                continue
            cell = row.cells[index]
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Solverstreet Technologies | BA Project Documentation"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Confidential client delivery artifact"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def para(doc, text="", style=None, bold=False, color=None, size=None, align=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return p


def bullet(doc, text):
    return para(doc, text, style="List Bullet")


def number(doc, text):
    return para(doc, text, style="List Number")


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    set_borders(tbl)
    set_table_widths(tbl, widths)
    hdr = tbl.rows[0].cells
    for i, value in enumerate(headers):
        hdr[i].text = value
        set_cell_shading(hdr[i], LIGHT_FILL)
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = NAVY
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
    set_table_widths(tbl, widths)
    para(doc, "", after=3)
    return tbl


def callout(doc, label, text):
    tbl = doc.add_table(rows=1, cols=1)
    set_borders(tbl)
    set_table_widths(tbl, [9360])
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = NAVY
    p.add_run(text)
    para(doc, "", after=3)


def add_doc_control(doc):
    para(doc, "Document Control", "Heading 1")
    table(
        doc,
        ["Field", "Details"],
        [
            ("Project", "Solverstreet Technologies AI and Digital Solutions Website"),
            ("Prepared By", "Business Analyst / Product Owner"),
            ("Document Type", "Integrated BA Pack: BRD, FRD, SRS, CRD, RTM, UAT, Handover"),
            ("Version", "1.0"),
            ("Date", date.today().strftime("%B %d, %Y")),
            ("Status", "Completed for client demonstration and delivery evidence"),
        ],
        [2200, 7160],
    )
    table(
        doc,
        ["Version", "Date", "Owner", "Description"],
        [
            ("0.1", "Discovery", "BA", "Captured client goals, website scope, user journeys, and technology expectations."),
            ("0.5", "Build Phase", "BA / Delivery", "Validated core pages, forms, authentication, dashboard, database, and setup requirements."),
            ("1.0", date.today().strftime("%b %d, %Y"), "BA / Owner", "Prepared final documentation after completion and production build verification."),
        ],
        [1200, 1700, 1900, 4560],
    )


def add_cover(doc):
    para(doc, "", after=24)
    p = para(doc, "BUSINESS ANALYSIS DOCUMENTATION PACK", size=12, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    p.paragraph_format.space_before = Pt(48)
    para(doc, "Solverstreet Technologies", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, "AI and Digital Solutions Company Website", size=16, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    para(doc, "BRD | FRD | SRS | Change Requirements | UAT | Handover", size=11, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    callout(
        doc,
        "Delivery Context",
        "This document represents the BA-owned requirements, scope, acceptance criteria, and project handover artifacts for a completed enterprise-grade website prototype."
    )
    para(doc, "Prepared as client-facing evidence of discovery, analysis, requirements management, validation, and delivery ownership.", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    doc.add_page_break()


def add_executive_summary(doc):
    para(doc, "Executive Summary", "Heading 1")
    para(doc, "Solverstreet Technologies required a premium, responsive, enterprise-ready website prototype to represent its AI, software engineering, automation, cloud, Salesforce, analytics, and digital transformation services. The completed solution provides a modern public website, contact inquiry flow, authentication, protected dashboard, Prisma data model, API routes, and setup documentation.")
    callout(doc, "Business Outcome", "The project positions Solverstreet Technologies as a credible AI and digital solutions provider and gives the business a maintainable foundation for sales, lead capture, client demos, and future production extension.")
    table(
        doc,
        ["Objective", "Delivered Evidence"],
        [
            ("Professional brand presence", "Responsive homepage, navigation, footer, hero, services, industries, portfolio, testimonials, pricing, blog, and legal pages."),
            ("Lead generation", "Validated inquiry form with service, budget, timeline, industry, company, and project description fields."),
            ("User account capability", "Register, login, logout, forgot/reset password, email verification route, and protected dashboard."),
            ("Scalable implementation", "Next.js App Router, React, TypeScript, Tailwind CSS, Prisma, PostgreSQL, NextAuth, Zod, and Nodemailer-ready mail hooks."),
            ("Delivery readiness", "Project builds successfully with 19 routes generated and TypeScript validation passing."),
        ],
        [3000, 6360],
    )


def add_ba_role(doc):
    para(doc, "BA Engagement Narrative", "Heading 1")
    para(doc, "As the Business Analyst and acting product owner, the engagement was handled from discovery through delivery validation. The BA role translated client expectations into structured requirements, clarified functional flows, prioritized scope, defined acceptance criteria, monitored change impact, and prepared final handover documentation.")
    for item in [
        "Conducted requirement discovery for brand, audience, service categories, lead capture, authentication, and dashboard expectations.",
        "Converted business goals into BRD, FRD, SRS, traceability, UAT, and change-control artifacts.",
        "Defined user journeys for visitors, prospects, registered users, administrators, and delivery stakeholders.",
        "Validated that the website reflects client positioning as an AI-first digital solutions company.",
        "Confirmed final technical verification through successful TypeScript and production build checks.",
    ]:
        bullet(doc, item)


def add_brd(doc):
    para(doc, "Business Requirements Document", "Heading 1")
    para(doc, "Business Need", "Heading 2")
    para(doc, "Solverstreet Technologies needs a polished digital front door that communicates its service capabilities, captures qualified leads, supports account registration, and provides a foundation that can evolve into a production business platform.")
    para(doc, "Business Objectives", "Heading 2")
    for item in [
        "Increase trust with enterprise prospects through a premium, credible website experience.",
        "Explain AI, automation, web development, Salesforce, cloud, and analytics services clearly.",
        "Capture structured project inquiries for sales qualification.",
        "Allow users to register, log in, and access a protected dashboard.",
        "Create an extensible codebase that future developers can maintain without redesigning the foundation.",
    ]:
        bullet(doc, item)
    para(doc, "Stakeholders", "Heading 2")
    table(
        doc,
        ["Stakeholder", "Interest", "Success Measure"],
        [
            ("Business Owner", "Brand credibility, lead generation, scalability", "Website supports demos and prospect conversion."),
            ("Prospective Client", "Clear service discovery and easy consultation request", "Can understand offerings and submit inquiry quickly."),
            ("Registered User", "Access to profile, inquiries, saved projects", "Can authenticate and view dashboard information."),
            ("Sales Team", "Structured inquiry details", "Receives useful service, budget, timeline, and contact data."),
            ("Development Team", "Maintainable architecture", "Reusable components, clear folders, typed validations, documented setup."),
        ],
        [2100, 4060, 3200],
    )
    para(doc, "Scope", "Heading 2")
    table(
        doc,
        ["In Scope", "Out of Scope / Future Phase"],
        [
            ("Public marketing website with responsive sections", "Live payment processing or quote generation"),
            ("Portfolio filtering and sample project content", "Real CMS-backed blog publishing workflow"),
            ("Inquiry form validation and database persistence", "Production object storage for uploaded files"),
            ("Credential authentication and protected dashboard", "Enterprise SSO and role-based admin console"),
            ("Prisma schema and PostgreSQL-ready backend", "Full production email template design and deliverability setup"),
        ],
        [4680, 4680],
    )


def add_frd(doc):
    para(doc, "Functional Requirements Document", "Heading 1")
    table(
        doc,
        ["ID", "Functional Requirement", "Priority", "Acceptance Criteria"],
        [
            ("FR-01", "Display sticky responsive header with Home, About, Services, Solutions, Industries, Portfolio, Pricing, Blog, and Contact navigation.", "Must", "Navigation is visible on desktop, collapses on mobile, and links to expected routes or sections."),
            ("FR-02", "Show animated hero with primary messaging and calls to action.", "Must", "Hero includes headline, subtitle, Book Consultation, and Explore Services actions."),
            ("FR-03", "Present company mission, values, service catalog, industry coverage, portfolio items, testimonials, and pricing preview.", "Must", "Each content area is visible, responsive, and aligned with Solverstreet Technologies positioning."),
            ("FR-04", "Allow visitors to submit structured project inquiries.", "Must", "Form validates required fields, sends request to API, stores inquiry, and shows success/error feedback."),
            ("FR-05", "Support registration, login, logout, forgot password, reset password, and email verification route.", "Must", "Auth pages and API routes exist and are connected to user/session models."),
            ("FR-06", "Protect dashboard content from unauthenticated users.", "Must", "Unauthenticated visitors are redirected to login; authenticated users see profile, inquiries, saved projects, and settings."),
            ("FR-07", "Support dark/light mode.", "Should", "Theme toggle switches the visual theme without breaking layout."),
            ("FR-08", "Provide legal, blog, and pricing pages for a complete website prototype.", "Should", "Routes render without build errors."),
        ],
        [900, 4300, 1100, 3060],
    )


def add_srs(doc):
    para(doc, "Software Requirements Specification", "Heading 1")
    para(doc, "System Overview", "Heading 2")
    para(doc, "The application is a Next.js App Router project using React, TypeScript, Tailwind CSS, Framer Motion, shadcn-style UI primitives, NextAuth.js, Prisma, PostgreSQL, React Hook Form, Zod, and Nodemailer-ready email integration.")
    para(doc, "Architecture Requirements", "Heading 2")
    for item in [
        "Use component-based structure under app, components, config, features, lib, types, and prisma folders.",
        "Use server components where appropriate and client components for interactive UI, forms, theme switching, and animation.",
        "Use Prisma as the database access layer and PostgreSQL as the target relational database.",
        "Use Zod schemas for form and API validation.",
        "Keep environment-specific secrets in .env and document required variables in .env.example.",
    ]:
        bullet(doc, item)
    para(doc, "Data Model Summary", "Heading 2")
    table(
        doc,
        ["Entity", "Purpose"],
        [
            ("User", "Stores registered user profile and authentication-related fields."),
            ("Account / Session", "Supports NextAuth account/session persistence."),
            ("VerificationToken", "Supports email verification workflow."),
            ("PasswordResetToken", "Supports forgot/reset password workflow."),
            ("Inquiry", "Stores project inquiry details submitted by prospects."),
            ("SavedProject", "Stores user-saved portfolio or project references."),
        ],
        [2600, 6760],
    )
    para(doc, "API Requirements", "Heading 2")
    table(
        doc,
        ["Route", "Purpose"],
        [
            ("/api/auth/[...nextauth]", "NextAuth credential authentication and session handling."),
            ("/api/register", "Create new user accounts with validation."),
            ("/api/inquiries", "Accept and persist contact/project inquiry submissions."),
            ("/api/forgot-password", "Generate reset token for valid users."),
            ("/api/reset-password", "Validate reset token and update password."),
            ("/api/verify-email", "Verify user email through token route."),
            ("/api/saved-projects", "Manage saved project data for authenticated users."),
        ],
        [2700, 6660],
    )
    para(doc, "Non-Functional Requirements", "Heading 2")
    table(
        doc,
        ["Category", "Requirement"],
        [
            ("Performance", "Pages should be optimized for fast load, code splitting, lazy loading, and successful production build."),
            ("Accessibility", "Use semantic sections, accessible buttons, labels, focus states, and readable contrast."),
            ("Security", "Protect dashboard routes, validate inputs, hash passwords, and keep secrets outside source code."),
            ("Maintainability", "Use reusable components, typed validation, isolated feature modules, and documented setup."),
            ("Responsiveness", "Support desktop, tablet, and mobile layouts without broken navigation or overflow."),
            ("SEO", "Provide metadata, descriptive page structure, and indexable public content."),
        ],
        [2200, 7160],
    )


def add_change_requirements(doc):
    para(doc, "Change Requirements and Control Log", "Heading 1")
    para(doc, "The following change records demonstrate BA ownership of scope clarification, impact assessment, and final project alignment.")
    table(
        doc,
        ["CR ID", "Change Request", "Impact", "Status"],
        [
            ("CR-01", "Create a complete enterprise-grade project folder instead of a single-page prototype.", "Expanded architecture, docs, app routes, backend routes, and maintainability requirements.", "Approved / Delivered"),
            ("CR-02", "Add authentication and protected dashboard.", "Added user flows, session handling, protected route, and database models.", "Approved / Delivered"),
            ("CR-03", "Add inquiry persistence and email notification hooks.", "Added Prisma Inquiry model, API validation, and Nodemailer-ready service.", "Approved / Delivered"),
            ("CR-04", "Remove external font dependency for offline build reliability.", "Improved local build stability by using system font stack.", "Approved / Delivered"),
            ("CR-05", "Finalize company brand as Solverstreet Technologies.", "Updated brand references across UI, metadata, mail, footer, auth copy, environment examples, and README.", "Approved / Delivered"),
        ],
        [1000, 3300, 3400, 1660],
    )
    para(doc, "Change Approval Criteria", "Heading 2")
    for item in [
        "Business value is clear and aligned with website positioning or delivery readiness.",
        "Technical impact is understood before implementation.",
        "Scope changes are reflected in relevant documentation and acceptance criteria.",
        "Final change is verified through build or functional checks where applicable.",
    ]:
        bullet(doc, item)


def add_user_stories(doc):
    para(doc, "User Stories and Acceptance Criteria", "Heading 1")
    table(
        doc,
        ["Story ID", "User Story", "Acceptance Criteria"],
        [
            ("US-01", "As a visitor, I want to understand Solverstreet Technologies services so that I can evaluate fit.", "Services, industries, portfolio, and value proposition are visible and easy to scan."),
            ("US-02", "As a prospect, I want to request a consultation so that I can start a project discussion.", "Contact form captures required details, validates input, and returns confirmation feedback."),
            ("US-03", "As a new user, I want to register so that I can access personalized dashboard features.", "Register page creates account through validated API route."),
            ("US-04", "As a returning user, I want to log in so that I can access my inquiries and saved projects.", "Login succeeds with valid credentials and dashboard is protected."),
            ("US-05", "As a business owner, I want qualified inquiry details so that sales follow-up is efficient.", "Inquiry includes company, service, budget, timeline, industry, and description fields."),
            ("US-06", "As a developer, I want clear setup instructions so that the project can be run and extended.", "README and environment template describe install, database, Prisma, and run steps."),
        ],
        [1000, 4200, 4160],
    )


def add_rtm_uat(doc):
    para(doc, "Requirements Traceability Matrix", "Heading 1")
    table(
        doc,
        ["Business Req", "Functional Req", "Implementation Area", "Verification"],
        [
            ("BR-01 Brand credibility", "FR-01, FR-02, FR-03", "Header, homepage, sections, footer", "Production build successful; visual routes generated."),
            ("BR-02 Lead capture", "FR-04", "Contact form, /api/inquiries, Prisma Inquiry", "Validation and API route implemented."),
            ("BR-03 User accounts", "FR-05, FR-06", "Auth pages, NextAuth route, dashboard", "TypeScript and build checks passed."),
            ("BR-04 Maintainability", "FR-07, FR-08", "src/features, src/components, src/lib, prisma", "Structured folder architecture delivered."),
            ("BR-05 Delivery readiness", "All FRs", "Next.js production build", "19 routes generated successfully."),
        ],
        [1900, 2100, 3000, 2360],
    )
    para(doc, "UAT Plan", "Heading 1")
    table(
        doc,
        ["Test Case", "Scenario", "Expected Result", "Status"],
        [
            ("UAT-01", "Open homepage on desktop and mobile viewport.", "Hero, navigation, sections, and footer render responsively.", "Ready for client sign-off"),
            ("UAT-02", "Use navigation links.", "Each route or page section opens as expected.", "Ready for client sign-off"),
            ("UAT-03", "Submit incomplete inquiry form.", "Validation errors display for missing or invalid fields.", "Ready for client sign-off"),
            ("UAT-04", "Submit complete inquiry form with database configured.", "Inquiry persists and success feedback is shown.", "Ready for client sign-off"),
            ("UAT-05", "Register, login, and open dashboard.", "Authenticated user reaches protected dashboard.", "Ready for client sign-off"),
            ("UAT-06", "Run production build.", "Build completes and generates all routes.", "Passed"),
        ],
        [1300, 3100, 3300, 1660],
    )


def add_handover(doc):
    para(doc, "Deployment and Handover Notes", "Heading 1")
    para(doc, "Local Setup", "Heading 2")
    for item in [
        "Install Node.js LTS and pnpm.",
        "Create .env from .env.example and configure DATABASE_URL, AUTH_SECRET, AUTH_URL, and SMTP values.",
        "Create PostgreSQL database ai_digital_solutions.",
        "Run pnpm install, pnpm run prisma:generate, pnpm run prisma:migrate, and pnpm run dev.",
        "Use pnpm run build for production validation before deployment.",
    ]:
        number(doc, item)
    para(doc, "Production Readiness Checklist", "Heading 2")
    for item in [
        "Replace placeholder images, sample portfolio items, legal copy, and example email addresses.",
        "Configure production PostgreSQL database and run migrations.",
        "Add real SMTP or transactional email provider.",
        "Add production file storage for uploaded documents.",
        "Review accessibility, SEO metadata, analytics, cookie policy, and security headers before launch.",
    ]:
        bullet(doc, item)
    para(doc, "Project Sign-Off", "Heading 1")
    table(
        doc,
        ["Role", "Name", "Sign-Off", "Date"],
        [
            ("Client Sponsor", "", "", ""),
            ("Business Analyst / Product Owner", "", "", ""),
            ("Technical Lead", "", "", ""),
            ("QA / UAT Reviewer", "", "", ""),
        ],
        [2400, 2400, 2400, 2160],
    )
    callout(doc, "Final BA Statement", "The documented scope, functional flows, technical requirements, change records, and validation evidence indicate that the Solverstreet Technologies website project is complete for demonstration and handover, with production environment configuration remaining as an operational deployment activity.")


def build():
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_doc_control(doc)
    add_executive_summary(doc)
    add_ba_role(doc)
    add_brd(doc)
    doc.add_page_break()
    add_frd(doc)
    doc.add_page_break()
    add_srs(doc)
    doc.add_page_break()
    add_change_requirements(doc)
    add_user_stories(doc)
    doc.add_page_break()
    add_rtm_uat(doc)
    doc.add_page_break()
    add_handover(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
